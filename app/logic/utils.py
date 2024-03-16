from docx import Document

def seconds_to_hms(seconds):
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60
    return "{:02d}:{:02d}:{:02d}".format(int(hours), int(minutes), int(seconds))

def create_doc(data):
    doc = Document()
    for d in data['data']:
        doc.add_heading(f"{d['speaker']}  {seconds_to_hms(d['start'])} - {seconds_to_hms(d['end'])}")
        doc.add_paragraph(f"{d['text']}")

    return doc